from __future__ import annotations

import json
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

try:
    import yaml  # type: ignore
except Exception as e:
    raise RuntimeError("PyYAML required. Install with: pip install pyyaml") from e

try:
    import jsonschema  # type: ignore
except Exception as e:
    raise RuntimeError("jsonschema required. Install with: pip install jsonschema") from e

from docx import Document
from docx.shared import Pt
from jinja2 import Environment, StrictUndefined, TemplateError
from jinja2.sandbox import SandboxedEnvironment

TEMPLATES_DIR = Path("templates")
CUSTOM_TEMPLATES_DIR = Path("templates/custom")
OUTPUT_DIR = Path("runs/ui/templates")
SCHEMA_PATH = Path("schemas/template.json")
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


class TemplateValidationError(Exception):
    """Raised when template validation fails with user-friendly message."""

    def __init__(self, template_path: str, errors: list[str]):
        self.template_path = template_path
        self.errors = errors
        message = f"Template validation failed for {template_path}:\n" + "\n".join(f"  - {e}" for e in errors)
        super().__init__(message)


class TemplateRenderError(Exception):
    """Raised when template rendering fails with user-friendly message."""

    pass


@dataclass
class InputDef:
    """Definition of a template input field."""

    id: str
    label: str
    type: str
    required: bool = False
    default: Any = None
    help: str = ""
    placeholder: str = ""
    validators: dict[str, Any] = field(default_factory=dict)


@dataclass
class TemplateDef:
    """Complete template definition with metadata and inputs."""

    path: Path
    name: str
    version: str
    description: str
    context: str
    inputs: list[InputDef]
    body: str
    style: str | None = None

    @property
    def key(self) -> str:
        """Generate key from filename."""
        return self.path.stem


def _load_schema() -> dict[str, Any]:
    """Load the template JSON schema."""
    if not SCHEMA_PATH.exists():
        raise FileNotFoundError(f"Template schema not found at {SCHEMA_PATH}")
    return json.loads(SCHEMA_PATH.read_text(encoding="utf-8"))


def _validate_template_data(data: dict[str, Any], template_path: str) -> None:
    """
    Validate template data against JSON schema.

    Args:
        data: Template YAML data
        template_path: Path to template file for error messages

    Raises:
        TemplateValidationError: If validation fails
    """
    schema = _load_schema()
    validator = jsonschema.Draft7Validator(schema)
    errors = []

    for error in validator.iter_errors(data):
        # Convert jsonschema errors to friendly messages
        field_path = ".".join(str(p) for p in error.path) if error.path else "root"
        errors.append(f"{field_path}: {error.message}")

    if errors:
        raise TemplateValidationError(template_path, errors)


def _load_yaml(path: Path) -> dict[str, Any]:
    """Load and parse YAML file."""
    try:
        return yaml.safe_load(path.read_text(encoding="utf-8")) or {}
    except yaml.YAMLError as e:
        raise TemplateValidationError(str(path), [f"Invalid YAML syntax: {e}"]) from e


def _parse_template(path: Path) -> TemplateDef:
    """
    Load and validate a template from YAML file.

    Args:
        path: Path to template YAML file

    Returns:
        TemplateDef object

    Raises:
        TemplateValidationError: If template is invalid
    """
    data = _load_yaml(path)

    # Validate against schema
    _validate_template_data(data, str(path))

    # Parse inputs
    inputs = []
    for inp_data in data.get("inputs", []):
        inputs.append(
            InputDef(
                id=inp_data["id"],
                label=inp_data["label"],
                type=inp_data["type"],
                required=inp_data.get("required", False),
                default=inp_data.get("default"),
                help=inp_data.get("help", ""),
                placeholder=inp_data.get("placeholder", ""),
                validators=inp_data.get("validators", {}),
            )
        )

    # Extract rendering body
    rendering = data.get("rendering", {})
    if not rendering.get("body"):
        raise TemplateValidationError(str(path), ["rendering.body is required"])

    return TemplateDef(
        path=path,
        name=data["name"],
        version=data["version"],
        description=data["description"],
        context=data["context"],
        inputs=inputs,
        body=rendering["body"],
        style=data.get("style"),
    )


def list_templates() -> list[TemplateDef]:
    """
    List all valid templates from templates/ directory.

    Returns:
        List of TemplateDef objects. Invalid templates are logged but not returned.
    """
    TEMPLATES_DIR.mkdir(exist_ok=True, parents=True)
    CUSTOM_TEMPLATES_DIR.mkdir(exist_ok=True, parents=True)

    out: list[TemplateDef] = []
    search_paths = [
        *sorted(TEMPLATES_DIR.glob("*.yaml")),
        *sorted(CUSTOM_TEMPLATES_DIR.glob("*.yaml")),
    ]

    for yml in search_paths:
        try:
            template = _parse_template(yml)
            out.append(template)
        except TemplateValidationError as e:
            # Log but don't crash - allow UI to show error
            print(f"Warning: {e}")
        except Exception as e:
            print(f"Warning: Failed to load {yml}: {e}")

    return out


# Custom Jinja2 filters for template safety
def to_slug(text: str) -> str:
    """Convert text to URL-safe slug."""
    text = str(text).lower().strip()
    text = re.sub(r"[^\w\s-]", "", text)
    text = re.sub(r"[-\s]+", "-", text)
    return text


def to_title(text: str) -> str:
    """Convert text to title case."""
    return str(text).title()


def _create_sandbox_env(context: str) -> Environment:
    """
    Create a sandboxed Jinja2 environment with safe defaults.

    Args:
        context: Output context ("markdown", "docx", "html")

    Returns:
        Configured SandboxedEnvironment
    """
    # Enable autoescape for HTML/DOCX contexts
    autoescape = context in ("html", "docx")

    env = SandboxedEnvironment(
        autoescape=autoescape,
        undefined=StrictUndefined,
        trim_blocks=True,
        lstrip_blocks=True,
    )

    # Add custom filters
    env.filters["to_slug"] = to_slug
    env.filters["to_title"] = to_title

    # Restrict to safe built-in filters only
    safe_filters = {
        "lower",
        "upper",
        "title",
        "replace",
        "join",
        "length",
        "round",
        "default",
        "safe",
        "escape",
        "e",
        "map",
        "select",
    }

    # Remove any unsafe filters
    for filter_name in list(env.filters.keys()):
        if filter_name not in safe_filters and not filter_name.startswith("to_"):
            del env.filters[filter_name]

    return env


def validate_inputs(template: TemplateDef, values: dict[str, Any]) -> list[str]:
    """
    Validate input values against template input definitions.

    Args:
        template: Template definition
        values: Input values to validate

    Returns:
        List of validation error messages (empty if valid)
    """
    errors = []

    for inp in template.inputs:
        value = values.get(inp.id)

        # Check required
        if inp.required and (value is None or value == ""):
            errors.append(f"{inp.label} is required")
            continue

        # Skip validation if empty and not required
        if value is None or value == "":
            continue

        # Type validation
        if inp.type == "int":
            try:
                val = int(value)
                if "min" in inp.validators and val < inp.validators["min"]:
                    errors.append(f"{inp.label} must be at least {inp.validators['min']}")
                if "max" in inp.validators and val > inp.validators["max"]:
                    errors.append(f"{inp.label} must be at most {inp.validators['max']}")
            except (ValueError, TypeError):
                errors.append(f"{inp.label} must be an integer")

        elif inp.type == "float":
            try:
                val = float(value)
                if "min" in inp.validators and val < inp.validators["min"]:
                    errors.append(f"{inp.label} must be at least {inp.validators['min']}")
                if "max" in inp.validators and val > inp.validators["max"]:
                    errors.append(f"{inp.label} must be at most {inp.validators['max']}")
            except (ValueError, TypeError):
                errors.append(f"{inp.label} must be a number")

        elif inp.type == "email":
            email_pattern = r"^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$"
            if not re.match(email_pattern, str(value)):
                errors.append(f"{inp.label} must be a valid email address")

        elif inp.type == "url":
            url_pattern = r"^https?://.+\..+"
            if not re.match(url_pattern, str(value)):
                errors.append(f"{inp.label} must be a valid URL (http:// or https://)")

        elif inp.type == "enum":
            choices = inp.validators.get("choices", [])
            if value not in choices:
                errors.append(f"{inp.label} must be one of: {', '.join(choices)}")

        elif inp.type == "multiselect":
            choices = inp.validators.get("choices", [])
            if isinstance(value, list):
                for v in value:
                    if v not in choices:
                        errors.append(f"{inp.label}: '{v}' is not a valid choice")

        # String length validation
        if inp.type in ("string", "text") and isinstance(value, str):
            if "min" in inp.validators and len(value) < inp.validators["min"]:
                errors.append(f"{inp.label} must be at least {inp.validators['min']} characters")
            if "max" in inp.validators and len(value) > inp.validators["max"]:
                errors.append(f"{inp.label} must be at most {inp.validators['max']} characters")

        # Regex validation
        if "regex" in inp.validators and isinstance(value, str):
            pattern = inp.validators["regex"]
            if not re.match(pattern, value):
                errors.append(f"{inp.label} does not match required pattern")

    return errors


def render_template(template: TemplateDef, variables: dict[str, Any]) -> str:
    """
    Render a template with sandboxed Jinja2 environment.

    Args:
        template: Template definition
        variables: Variable values for rendering

    Returns:
        Rendered template text

    Raises:
        TemplateRenderError: If rendering fails
    """
    # Validate inputs first
    validation_errors = validate_inputs(template, variables)
    if validation_errors:
        raise TemplateRenderError("Validation failed:\n" + "\n".join(f"  - {e}" for e in validation_errors))

    # Create sandboxed environment
    env = _create_sandbox_env(template.context)

    try:
        tmpl = env.from_string(template.body)
        return tmpl.render(**variables)
    except TemplateError as e:
        raise TemplateRenderError(f"Template rendering failed: {e}") from e
    except Exception as e:
        raise TemplateRenderError(f"Unexpected error during rendering: {e}") from e


def export_markdown(text: str, fname: str) -> Path:
    """
    Export text as markdown file.

    Args:
        text: Content to export
        fname: Base filename (without extension)

    Returns:
        Path to created file
    """
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    p = OUTPUT_DIR / f"{fname}.md"
    p.write_text(text, encoding="utf-8")
    return p


def export_docx(text: str, fname: str, heading: str = "DJP Output", style_path: str | None = None) -> Path:
    """
    Export text as DOCX file.

    Args:
        text: Content to export
        fname: Base filename (without extension)
        heading: Document heading
        style_path: Optional path to base style DOCX (deferred for Sprint 2)

    Returns:
        Path to created file
    """
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    # Note: style_path support deferred to Sprint 2
    doc = Document()
    doc.add_heading(heading, level=1)

    for p in doc.paragraphs:
        for run in p.runs:
            run.font.size = Pt(12)

    for line in text.splitlines():
        para = doc.add_paragraph(line)
        for run in para.runs:
            run.font.size = Pt(11)

    p = OUTPUT_DIR / f"{fname}.docx"
    doc.save(p)
    return p
